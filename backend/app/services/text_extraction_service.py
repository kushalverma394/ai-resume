from __future__ import annotations

import logging
from pathlib import Path

import pdfplumber
from docx import Document
from fastapi import HTTPException, status

from app.schemas.resume import ExtractedResume
from app.utils.file_utils import make_bytes_io, preview_text, validate_resume_file


logger = logging.getLogger(__name__)


class TextExtractionService:
    def extract(self, filename: str, content_type: str, data: bytes, size_bytes: int) -> ExtractedResume:
        validate_resume_file(filename, content_type)
        suffix = Path(filename).suffix.lower()

        if suffix == ".pdf":
            text = self._extract_pdf_text(data)
        elif suffix == ".docx":
            text = self._extract_docx_text(data)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Only PDF and DOCX files are supported.",
            )

        cleaned_text = self._clean_text(text)
        if not cleaned_text:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="No readable text could be extracted from the uploaded file.",
            )

        logger.info("Extracted %s characters from %s", len(cleaned_text), filename)
        return ExtractedResume(
            filename=filename,
            content_type=content_type,
            size_bytes=size_bytes,
            text=cleaned_text,
            text_preview=preview_text(cleaned_text),
        )

    def _extract_pdf_text(self, data: bytes) -> str:
        try:
            with pdfplumber.open(make_bytes_io(data)) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Unable to read PDF content: {exc}",
            ) from exc

    def _extract_docx_text(self, data: bytes) -> str:
        try:
            document = Document(make_bytes_io(data))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_text: list[str] = []
            for table in document.tables:
                for row in table.rows:
                    row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_values:
                        table_text.append(" ".join(row_values))
            return "\n".join(paragraphs + table_text)
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Unable to read DOCX content: {exc}",
            ) from exc

    @staticmethod
    def _clean_text(text: str) -> str:
        return "\n".join(line.strip() for line in text.splitlines() if line.strip())